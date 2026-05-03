import re
import PyPDF2
import docx
from typing import Dict, List, Any, Optional
import os

class TextExtractor:
    """Utility class for extracting and processing text from various document formats"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and metadata from PDF file
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            text = ""
            metadata = {
                "page_count": 0,
                "has_images": False,
                "has_tables": False,
                "file_size": os.path.getsize(file_path),
                "file_name": os.path.basename(file_path)
            }
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["page_count"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    
                    # Check for images (simplified)
                    if '/XObject' in page['/Resources']:
                        metadata["has_images"] = True
                    
                    # Check for tables (simplified)
                    if 'Table' in page_text or 'table' in page_text:
                        metadata["has_tables"] = True
            
            return {
                "success": True,
                "text": self._clean_text(text),
                "metadata": metadata,
                "word_count": len(text.split()),
                "character_count": len(text)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "text": ""
            }
    
    def extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and metadata from DOCX file
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            doc = docx.Document(file_path)
            text = ""
            metadata = {
                "page_count": len(doc.paragraphs) // 30,  # Rough estimate
                "has_images": len(doc.inline_shapes) > 0,
                "has_tables": len(doc.tables) > 0,
                "file_size": os.path.getsize(file_path),
                "file_name": os.path.basename(file_path)
            }
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        row_text += cell.text + " | "
                    text += row_text + "\n"
            
            return {
                "success": True,
                "text": self._clean_text(text),
                "metadata": metadata,
                "word_count": len(text.split()),
                "character_count": len(text)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "text": ""
            }
    
    def extract_from_txt(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from plain text file
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                "success": True,
                "text": self._clean_text(text),
                "metadata": {
                    "file_size": os.path.getsize(file_path),
                    "file_name": os.path.basename(file_path)
                },
                "word_count": len(text.split()),
                "character_count": len(text)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "text": ""
            }
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing extra whitespace and special characters
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters (keep basic punctuation)
        text = re.sub(r'[^\w\s\.\,\!\?\-\'\"]', '', text)
        
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        
        # Trim leading/trailing spaces
        text = text.strip()
        
        return text
    
    def split_into_sections(self, text: str) -> List[Dict[str, str]]:
        """
        Split text into logical sections based on headings
        
        Args:
            text: Extracted text
            
        Returns:
            List of sections with title and content
        """
        sections = []
        
        # Common heading patterns
        heading_patterns = [
            r'^(Chapter|Section|Part)\s+\d+[\.:]\s+(.+)$',
            r'^(\d+\.\s+[A-Z][A-Za-z\s]+)$',
            r'^([A-Z][A-Za-z\s]+:)$',
            r'^(Introduction|Background|Methodology|Results|Discussion|Conclusion)$'
        ]
        
        lines = text.split('\n')
        current_section = {"title": "Introduction", "content": []}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            is_heading = False
            for pattern in heading_patterns:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    # Save previous section
                    if current_section["content"]:
                        current_section["content"] = " ".join(current_section["content"])
                        sections.append(current_section.copy())
                    
                    # Start new section
                    current_section = {
                        "title": line,
                        "content": []
                    }
                    is_heading = True
                    break
            
            if not is_heading:
                current_section["content"].append(line)
        
        # Add last section
        if current_section["content"]:
            current_section["content"] = " ".join(current_section["content"])
            sections.append(current_section)
        
        return sections
    
    def extract_key_terms(self, text: str, top_n: int = 20) -> List[Dict[str, Any]]:
        """
        Extract key terms and their frequencies from text
        
        Args:
            text: Input text
            top_n: Number of top terms to return
            
        Returns:
            List of key terms with frequencies
        """
        # Stop words to ignore
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
            'have', 'has', 'had', 'having', 'do', 'does', 'did', 'doing', 'this',
            'that', 'these', 'those', 'there', 'their', 'they', 'we', 'you', 'he',
            'she', 'it', 'from', 'as', 'will', 'can', 'could', 'would', 'should'
        }
        
        # Split into words and convert to lowercase
        words = re.findall(r'\b[a-z]+(?:[a-z]+)?\b', text.lower())
        
        # Count word frequencies
        word_freq = {}
        for word in words:
            if len(word) > 3 and word not in stop_words:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Sort by frequency and get top N
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        
        return [
            {"term": word, "frequency": freq, "weight": freq / len(words)}
            for word, freq in sorted_words[:top_n]
        ]
    
    def detect_language(self, text: str) -> str:
        """
        Detect the language of the text (simplified)
        
        Args:
            text: Input text
            
        Returns:
            Detected language code
        """
        # Simple language detection based on common words
        # In production, use a proper library like langdetect
        common_english = {'the', 'and', 'to', 'of', 'a', 'in', 'is', 'it', 'you', 'that'}
        common_urdu = {'ہے', 'اور', 'کے', 'میں', 'نے', 'کا', 'کی', 'یہ', 'وہ', 'تو'}
        
        words = set(re.findall(r'\b[a-zA-Z]+\b', text.lower()))
        english_count = len(words.intersection(common_english))
        urdu_count = len(words.intersection(common_urdu))
        
        if urdu_count > english_count:
            return "ur"
        return "en"